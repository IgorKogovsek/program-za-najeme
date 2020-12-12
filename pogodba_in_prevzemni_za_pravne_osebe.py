import datetime
import tkinter as tk
from tkinter import *
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

suffix = datetime.datetime.now().strftime("%y%m%d_%H%M%S")

def kreiraj_pogodbo_in_zapisnik_PO():

    # najemna pogodba

    global master

    print(e1.get(), e2.get(), e3.get(), e4.get(), e5.get(), e6.get(), e7.get(), e8.get(), zastopnik.get(), e10.get(), e11.get(),
          e12.get(), e13.get(), e14.get(), e15.get(), e16.get(), e17.get(), e18.get(), lokacija.get(), skrbnik.get(), fakturiranje.get())

    input_ime_podjetja = e1.get()

    input_ulica = e2.get()

    input_hisna_stevilka = e3.get()

    input_postna_stevilka = e4.get()

    input_mesto = e5.get()

    input_drzava = e6.get()

    input_maticna_stevilka = e7.get()

    input_davcna_stevilka = e8.get()

    input_zastopnik = zastopnik.get()

    input_ime_zastopnika = e10.get()

    input_znamka = e11.get()

    input_model = e12.get()

    input_vin = e13.get()

    input_reg = e14.get()

    input_reg_do = e15.get()

    input_najem_od = e16.get()

    input_najem_do = e17.get()

    input_obrok_brez_DDV = e18.get()

    input_fakturiranje = fakturiranje.get()

    input_lokacija = lokacija.get()

    input_skrbnik = skrbnik.get()


    doc = docx.Document()

    najemodajalec = doc.add_paragraph(
        "SELMAR, d.o.o., Mariborska cesta 119, 3000 Celje, matična številka: 5300541, ID številka za DDV: SI62243217, transakcijski račun pri ABANKA d.d.: SI56 0510 0801 4707 908 , ki jo zastopa direktor: Uroš Seles (v nadaljevanju: najemodajalec) \n"
        "\n")
    najemodajalec.add_run("in")

    najemojemalec = doc.add_paragraph()
    najemojemalec.add_run(
        input_ime_podjetja + ", " + input_ulica + " " + input_hisna_stevilka + ", " + input_postna_stevilka + " " + input_mesto + ", " + input_drzava + ", matična številka: " + input_maticna_stevilka + ", ID številka za DDV: " + input_davcna_stevilka + ", ki jo zastopa " + input_zastopnik + " " + input_ime_zastopnika + " (v nadaljevanju: najemojemalec) \n")

    najemojemalec.add_run("\n")
    najemojemalec.add_run("sklepata")

    doc.add_heading('NAJEMNO POGODBO \n', level=2)

    prvi_clen = doc.add_paragraph()
    prvi_clen.add_run("1. ČAS TRAJANJA POGODBE \n").bold = True
    prvi_clen.add_run(
        "1. Najemna pogodba velja za obdobje od vključno " + input_najem_od + " do vključno " + input_najem_do + ", z možnostjo podaljšanja. V primeru podaljšanja se sklene nova pogodba. \n")

    drugi_clen = doc.add_paragraph()
    drugi_clen.add_run("2. PREDMET NAJEMNE POGODBE \n").bold = True
    drugi_clen.add_run(
        "1. Predmet najema pogodbe je " + input_znamka + " " + input_model + ", št. šasije: " + input_vin + ", reg. oznaka: " + input_reg + ", za obdobje, dogovorjeno v 1. členu te pogodbe. \n")

    tretji_clen = doc.add_paragraph()
    tretji_clen.add_run("3. CENA NAJEMA \n").bold = True
    tretji_clen.add_run(
        "1. Cena najema znaša " + input_obrok_brez_DDV + "EUR/" + input_fakturiranje + " za obdobje iz 1. člena najemne pogodbe in se plačuje mesečno na TRR podjetja Selmar, d.o.o. \n"
                                                                 "2. V ceno je vključeno zavarovanje, priprava vozila, dovoljenje za vožnjo izven Slovenije (če država namembnosti ni članica Evropske Unije). V ceno ni vključen 22% DDV in gorivo. \n"
                                                                 "3. Vsi morebitni prometni prekrški, ki nastanejo v času najema, bremenijo najemnika. \n")

    cetrti_clen = doc.add_paragraph()
    cetrti_clen.add_run("4. PREVZEM, VRNITEV, ČIŠČENJE \n").bold = True
    cetrti_clen.add_run("1. Vozilo se prevzame po 8.00 uri na dan začetka najemnega obdobja. \n"
                        "2. Vrnitev vozila se izvede zadnji dan najema do 17.00 ure. \n"
                        "3. Podaljšanje najema je skladno z dogovorom iz 1. člena te pogodbe. Prevzem in predaja vozila se vrši na sedežu podjetja na Mariborski cesti 119 v Celju ali po dogovoru. \n"
                        "4. Vozilo mora biti pri vrnitvi očiščeno. \n"
                        "5. Če najemnik vrne vozilo neočiščeno, je dolžan plačati najemodajalcu: \n"
                        "- 40,00 EUR za zunanje čiščenje, \n"
                        "- 40,00 EUR za notranje čiščenje. \n"
                        "6. Ob prevzemu se pregleda celotno vozilo (oprema, mehanika, karoserija, motor itd.), preveri se delovanje naprav v vozilu in sestavi se prevzemni zapisnik. Vse pomanjkljivosti, manjkajoča oprema, poškodbe  itd. se vpišejo v prevzemni zapisnik. Vse ugotovljene poškodbe se fotografirajo in dokumentirajo. Ob vrnitvi se s primopredajnim zapisnikom ugotovi stanje vozila in pregleda delovanje vseh naprav. Najemnik odgovarja za vse pomanjkljivosti in poškodbe, ki niso bile ugotovljene ob prevzemu. \n"
                        "7. V primeru višje sile, ki je razlog zamude pri vračilu vozila, je dolžan najemnik telefonsko obvestiti najemodajalca za vzrok zamude in predviden čas vračila vozila. \n"
                        "8. Ta pogodba lahko preneha na podlagi sporazuma strank (sporazumno) ali na podlagi odstopa najemodajalca. \n"
                        "Sporazumno se pogodba prekine v primeru, da se stranki o prekinitvi dogovorita in skleneta posebni pisni sporazum ali dodatek k tej pogodbi. \n"
                        "V primeru sporazumne prekinitve pogodbe je najemnik dolžan: \n"
                        "•	vrniti vozilo oziroma omogočiti odjavo vozila iz prometa, \n"
                        "•	plačati nadomestilo za morebitno negospodarno uporabo vozila, \n"
                        "•	plačati vse zapadle oziroma neplačane obveznosti do najemodajalca, \n"
                        "•	plačano dogovorjeno, z najemnino nepokrito izgubo vrednosti vozila, \n"
                        "•	plačati stroške izterjave napravočasno plačanih obveznosti, \n"
                        "•	plačati stroške morebitnega odvzema vozila in morebitne druge stroške, nastale s predmetom pogodbe. \n"
                        "\n"
                        "Ne glede na dejstvo, da je pogodba sklenjena za določen čas, lahko v naslednjih primerih najemodajalec od pogodbe odstopi oziroma pogodbo brez odpovednega roka odpove: \n"
                        "•	v primeru, da najemnik več kot enaindvajset (21) dni zamuja s plačilom obveznosti po tej pogodbi, \n"
                        "•	v primeru, da z vozilom ne ravna kot dober gospodar ter ne skrbi za redno vzdrževanje in servisiranje vozila, \n"
                        "•	brez soglasja najemodajalca vgrajuje ali odstranjuje dodatno opremo, \n"
                        "•	v primeru, da najemodajalca ne obvesti o spremembi naslova sedeža najemnika, spremembi TRR, prodaji podjetja ali uvedbi insolvenčnega postopka, \n"
                        "•	v primeru, da ob sklenitvi pogodbe ni navedel točnih podatkov o svojem gospodarskem stanju, zamolčal dejstva in okoliščine, predložil krive listine, zaradi katerih bi najemodajalec, če bi zanje vedel, ne sklenil te pogodbe, \n"
                        "•	v primeru, da ovira oziroma ne omogoči pravočasne registracije vozila ali pravočasno ne predloži dokumentov, \n"
                        "•	v primeru, da najemnik kljub opominu najemodajalca vozilo uporablja v nasprotju z deklarirano uporabo ali če vozilu ali če vozilu grozi škoda zaradi najemnikove opustitve vzdrževanja vozila. \n"
                        "\n"
                        "Najemnik mora v primerih iz zgornjega odstavka najkasneje v treh (3) dneh najemniku vrniti vozilo z vso prevzeto dokumentacijo in kompletom ključev. \n"
                        "\n"
                        "V primerih prenehanja pogodbe po tem členu je najemnik, poleg dolžnosti iz člena te pogodbe, najemodajalcu dolžan izpolniti še naslednje: \n"
                        "•	plačati stroške izterjave nepravočasno plačanih obveznosti, \n"
                        "•	plačati stroške morebitnega odvzema vozila, \n"
                        "•	plačati odškodnino za prekinitev pogodbe v višini vseh obveznosti za celotno obdobje, ki je bilo določeno ob sklenitvi pogodbe. \n")

    peti_clen = doc.add_paragraph()
    peti_clen.add_run("5. VOZNIK \n").bold = True
    peti_clen.add_run(
        "1. Voznik (oziroma vozniki) vozila mora na dan prevzema vozila  imeti dopolnjenih 21 let in posedovati veljavno vozniško dovoljenje vsaj tri leta. \n"
        "2. Voznik (oziroma vozniki) se obveže, da pred in med vožnjo ne bo užival alkoholnih pijač oziroma ne bo vozil pod vplivom substanc, ki zmanjšujejo psihomotorične sposobnosti (zdravila, droge, ipd.). \n")

    sesti_clen = doc.add_paragraph()
    sesti_clen.add_run("6. POPRAVILA \n").bold = True
    sesti_clen.add_run(
        "1. V primeru okvare vozila v času najema poskuša najemnik odpraviti okvaro na najbližjem pooblaščenem servisu po predhodnem dogovoru z najemodajalcem oz. izključno na servisu najemodajalca. Stroški popravila v primeru malomarnega ravnanja najemnika so breme najemnika. \n")

    sedmi_clen = doc.add_paragraph()
    sedmi_clen.add_run("7. OBNAŠANJE PRI NEZGODI \n").bold = True
    sedmi_clen.add_run(
        "1. Najemnik se obveže, da bo morebitno škodo na vozilu pri kakršnikoli nesreči (prometna nesreča, kraja, vlom v vozilo, poškodba na parkirišču itd.) obvezno prijavil policiji (ki je pristojna za tisto državo, v kateri se je nesreča pripetila) in takoj poklical najemodajalca. \n"
        "2. Najemnik mora najemodajalcu izročiti ustrezno dokumentacijo in priložiti skico o poškodbah vozila. Nezgodni zapisnik mora vsebovati imena in podatke vseh udeležencev nezgode. \n"
        "3. V primeru prometne nezgode, ko je voznik vinjen, stroške popravila krije najemnik. \n")

    osmi_clen = doc.add_paragraph()
    osmi_clen.add_run("8. ZAVAROVANJE \n").bold = True
    osmi_clen.add_run("1. Vozilo je obvezno in kasko zavarovano. \n")

    deveti_clen = doc.add_paragraph()
    deveti_clen.add_run("9. ODGOVORNOST NAJEMNIKA \n").bold = True
    deveti_clen.add_run(
        "1. Z vozilom lahko upravlja izključno najemnik, ki je naveden v pogodbi. Oziroma osebe, ki so potrjene s strani najemodajalca ali najemojemalca. \n"
        "2. Najemnik mora upoštevati cestno prometne predpise, sicer lahko v primeru morebitne prometne nesreče nastopijo dodatni zapleti s prometno policijo in kasneje z zavarovalnico. \n"
        "3. Najemnik odgovarja v vrednosti 1% odbitne franšize  v sklopu kasko zavarovanja za morebitno škodo zaradi njegove krivde in v zvezi s tem najemodajalčevih stroškov izgube bonusa. \n"
        "4. Najemnik odgovarja neomejeno v primeru škode zaradi zaradi preseženega tovora in ostalih primerov nedovoljene uporabe. \n"
        "5. V primeru, da zavarovalnica iz kakršnihkoli vzrokov zavrne plačilo, je dolžan škodo pokriti najemnik. \n"
        "6. Najemnik se obveže, da dokumente in ključe vozila ob zapustitvi le-tega nosi s seboj. V nasprotnem primeru zavarovalnica ne krije nastalih stroškov v primeru kraje vozila, kar pomeni, da vse stroške krije najemnik sam. \n"
        "7. Najemnik mora z vozilom ravnati kot skrben in vesten gospodar. \n"
        "8. Najemnik lahko v času najema vozila opravi največ 5.000 km na mesec. V kolikor najemnik prevozi več kot 5.000 km na mesec se mu vsak dodatni kilometer obračuna po vrednosti EUR 0,37 na km brez DDV. \n"
        "9. Najemnik se obvezuje redno plačevati mesečno najemnino. V kolikor najemnik s plačilom zamuja za dve najemnine, se mu vozilo takoj odvzame brez predhodnega opominjanja ter se mu zaračunajo stroški odvzema vozila. \n")

    deseti_clen = doc.add_paragraph()
    deseti_clen.add_run("10. ODGOVORNOST NAJEMODAJALCA \n").bold = True
    deseti_clen.add_run(
        "1. Najemodajalec skrbi, da je vozilo v času primopredaje tehnično brezhibno, brez kakršnekoli okvare, ki bi lahko povzročila nezgodo. \n"
        "2. Najemodajalec ne odgovarja za stvari, ki so bile puščene ali pozabljene v vozilu. V primeru višje sile, poškodovanega ali nevoznega vozila se dogovori za drug termin najema vozila. \n"
        "3. V nobenem primeru podjetje SELMAR d.o.o. ne prevzema nobene druge odgovornosti. \n")

    enajsti_clen = doc.add_paragraph()
    enajsti_clen.add_run("11. PRISTOJNO SODIŠČE \n").bold = True
    enajsti_clen.add_run(
        "1. V primeru neupoštevanja pogodbenih členov ali drugih morebitnih sporov je pristojno Okrožno sodišče v Celju. \n")

    table = doc.add_table(rows=6, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Najemodajalec:"
    hdr_cells[3].text = "Najemojemalec:"

    datum_cells = table.rows[1].cells
    datum_cells[0].text = "Datum:"
    datum_cells[3].text = "Datum:"

    podjetje_cells = table.rows[3].cells
    podjetje_cells[0].text = "Selmar d.o.o."
    podjetje_cells[3].text = input_ime_podjetja

    zastopnik_ime_cells = table.rows[4].cells
    zastopnik_ime_cells[0].text = "Uroš Seles"
    zastopnik_ime_cells[3].text = input_ime_zastopnika

    zastopnik_cells = table.rows[5].cells
    zastopnik_cells[0].text = "direktor"
    zastopnik_cells[3].text = input_zastopnik

    pogodba_name = "_".join(["najemna_pogodba", input_model, input_vin, input_ime_podjetja, suffix])


#vnos podatkov excel
    df = pd.read_excel("seznam_najemov_new.xlsx")

    new_row = {"znamka": input_znamka,
               "model": input_model,
               "vin": input_vin,
               "reg. št." : input_reg,
               "reg. do" : input_reg_do,
               "podjetje" : input_ime_podjetja,
               "najem od" : input_najem_od,
               "najem do" : input_reg_do,
               "obrok brez DDV" : input_obrok_brez_DDV,
               "fakturirano" : input_fakturiranje,
               "ulica" : input_ulica,
               "hišna št." : input_hisna_stevilka,
               "poštna št." : input_postna_stevilka,
               "mesto" : input_mesto,
               "država" : input_drzava,
               "zastopnik podjetja" : input_ime_zastopnika,
               "PE" : input_lokacija,
               "skrbnik" : input_skrbnik,
               "datum in ura" : suffix}

    df = df.append(new_row, ignore_index=True)

    df.drop(df.columns[df.columns.str.contains('unnamed', case=False)], axis=1, inplace=True)


    # prevzemni zapisnik

    print(e1.get(), e2.get(), e3.get(), e4.get(), e5.get(), e6.get(), e7.get(), e8.get(), zastopnik.get(), e10.get(), e11.get(),
          e12.get(), e13.get(), e14.get(), e15.get(), e16.get(), e17.get(), e18.get(), lokacija.get(), skrbnik.get(), fakturiranje.get())

    input_ime_podjetja = e1.get()

    input_ulica = e2.get()

    input_hisna_stevilka = e3.get()

    input_postna_stevilka = e4.get()

    input_mesto = e5.get()

    input_drzava = e6.get()

    input_maticna_stevilka = e7.get()

    input_davcna_stevilka = e8.get()

    input_zastopnik = zastopnik.get()

    input_ime_zastopnika = e10.get()

    input_znamka = e11.get()

    input_model = e12.get()

    input_vin = e13.get()

    input_reg = e14.get()

    input_reg_do = e15.get()

    input_najem_od = e16.get()

    input_najem_do = e17.get()

    input_mesecni_obrok_brez_DDV = e18.get()

    input_fakturiranje = fakturiranje.get()

    input_lokacija = lokacija.get()

    input_skrbnik = skrbnik.get()

    doc1 = docx.Document()

    #GLAVA PREVZEMNEGA ZAPISNIKA

    header = doc1.sections[0].header
    htable = header.add_table(1, 1, Inches(6))
    htab_cells = htable.rows[0].cells
    ht0 = htab_cells[0].add_paragraph()
    ht0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    kh = ht0.add_run()
    kh.add_picture('SELMAR logo.gif', width=Inches(1))

    #PREVZEMNI ZAPISNIK - PREVZEM

    prevzemnizapisnik = doc1.add_heading(
        "PREVZEMNI ZAPISNIK - prevzem", level=2)

    predajalec_prevzemalec_in_lokacija= doc1.add_paragraph()
    predajalec_prevzemalec_in_lokacija.add_run(
        "VOZILO IZDAL: SELMAR, d.o.o., Mariborska cesta 119, 3000 Celje, Slovenija \n"
    )
    predajalec_prevzemalec_in_lokacija.add_run(
        "VOZILO PREJEL: " + input_ime_podjetja + ", " + input_ulica + " " + input_hisna_stevilka + ", " + input_postna_stevilka + " " + input_mesto + ", " + input_drzava + "\n"
    )
    predajalec_prevzemalec_in_lokacija.add_run(
        "LOKACIJA PREVZEMA: " + input_lokacija
    )

    podatki_o_vozilu = doc1.add_paragraph()
    podatki_o_vozilu.add_run(
        "PODATKI O VOZILU \n").bold = True
    podatki_o_vozilu.add_run(
        "ZNAMKA IN MODEL: " + input_znamka + " " + input_model + " ; " + "VIN: " + input_vin + " ; \n"
    )
    podatki_o_vozilu.add_run(
        "REGISTRSKA OZNAKA: " + input_reg + " ; " + "REGISTRIRAN DO: " + input_reg_do
    )

    oprema = doc1.add_paragraph()
    oprema.add_run(
        "OPREMA: \n"
    ).bold = True
    oprema.add_run(
        "Število ključev: __________________ ; Prometno dovoljenje: DA \ NE ; Homologacija: DA \ NE; \n"
        "Servisna knjiga: DA \ NE ; Navodila za uporabo: DA \ NE ; Trikotnik in prva pomoč: DA \ NE; \n"
        "Polnilni kabel (1-fazni): DA \ NE ; Polnilni kabel (3-fazni): DA \ NE")

    datum_in_ura_prevzema = doc1.add_paragraph()
    datum_in_ura_prevzema.add_run(
        "DATUM IN URA PREVZEMA: ______________________________"
    )

    poskodbe = doc1.add_paragraph()
    poskodbe.add_run(
        "POŠKODBE OB PREVZEMU: "
    ).bold = True
    doc1.add_picture("vehicle_check_sheet1.png", width=Inches(6.3))

    opombe = doc1.add_paragraph()
    opombe.add_run(
        "Opombe: \n"
    ).bold = True
    opombe.add_run("________________________________________________________________________________________________________ \n")
    opombe.add_run("________________________________________________________________________________________________________")
    opombe.add_run("\n" + "\n" + "\n")

    podpis = doc1.add_paragraph()
    podpis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    podpis.add_run("____________________                                                                                                         ____________________ \n")
    podpis.add_run("Selmar d.o.o.                                                                                                                         Prevzel")

    doc1.add_page_break()

    #PREVZEMNI ZAPISNIK - VRACILO

    prevzemnizapisnik = doc1.add_heading(
        "PREVZEMNI ZAPISNIK - vračilo", level=2)

    predajalec_prevzemalec_in_lokacija= doc1.add_paragraph()
    predajalec_prevzemalec_in_lokacija.add_run(
        "VOZILO VRNIL: " + input_ime_podjetja + ", " + input_ulica + " " + input_hisna_stevilka + ", " + input_postna_stevilka + " " + input_mesto + ", " + input_drzava + "\n"
    )
    predajalec_prevzemalec_in_lokacija.add_run(
        "VOZILO PREJEL: SELMAR, d.o.o., Mariborska cesta 119, 3000 Celje, Slovenija \n"
    )
    predajalec_prevzemalec_in_lokacija.add_run(
        "LOKACIJA VRAČILA: _________________"
    )

    podatki_o_vozilu = doc1.add_paragraph()
    podatki_o_vozilu.add_run(
        "PODATKI O VOZILU \n").bold = True
    podatki_o_vozilu.add_run(
        "ZNAMKA IN MODEL: " + input_znamka + " " + input_model + " ; " + "VIN: " + input_vin + " ; \n"
    )
    podatki_o_vozilu.add_run(
        "REGISTRSKA OZNAKA: " + input_reg + " ; " + "REGISTRIRAN DO: " + input_reg_do
    )

    oprema = doc1.add_paragraph()
    oprema.add_run(
        "OPREMA: \n"
    ).bold = True
    oprema.add_run(
        "Število ključev: __________________ ; Prometno dovoljenje: DA \ NE ; Homologacija: DA \ NE; \n"
        "Servisna knjiga: DA \ NE ; Navodila za uporabo: DA \ NE ; Trikotnik in prva pomoč: DA \ NE; \n"
        "Polnilni kabel (1-fazni): DA \ NE ; Polnilni kabel (3-fazni): DA \ NE")

    datum_in_ura_prevzema = doc1.add_paragraph()
    datum_in_ura_prevzema.add_run(
        "DATUM IN URA VRAČILA: ______________________________"
    )

    poskodbe = doc1.add_paragraph()
    poskodbe.add_run(
        "POŠKODBE OB VRAČILU: "
    ).bold = True
    doc1.add_picture("vehicle_check_sheet1.png", width=Inches(6.3))

    opombe = doc1.add_paragraph()
    opombe.add_run(
        "Opombe: \n"
    ).bold = True
    opombe.add_run("________________________________________________________________________________________________________ \n")
    opombe.add_run("________________________________________________________________________________________________________")
    opombe.add_run("\n" + "\n" + "\n")

    podpis = doc1.add_paragraph()
    podpis.alignment = WD_ALIGN_PARAGRAPH.CENTER
    podpis.add_run("____________________                                                                                                         ____________________ \n")
    podpis.add_run("Selmar d.o.o.                                                                                                                          Vrnil")

    prevzemni_name = "_".join(["prevzemni_zapisnik", input_model, input_vin, input_ime_podjetja, suffix])

    # preveri ali so vsi podatki izpolnjeni

    if e1.get() and e2.get() and e3.get() and e4.get() and e5.get() and e6.get() and e7.get() and e8.get() and zastopnik.get() != "Izberi" and e10.get() and e11.get() and e12.get() and e13.get() and e14.get() and e15.get() and e16.get() and e17.get() and e18.get() and fakturiranje.get() != "Izberi" and lokacija.get() != "Izberi" and skrbnik.get() != "Izberi":
        master3 = Tk()
        frame1 = Frame(master3, highlightbackground="green", highlightcolor="green", highlightthickness=1, bd=0)
        frame1.pack()
        master3.overrideredirect(1)
        master3.geometry("200x70+650+400")
        lbl = Label(frame1, text="Potrdi vnos")
        lbl.pack()
        yes_btn = Button(frame1, text="Da", bg="light blue", fg="red", command=lambda: [doc.save(pogodba_name + ".docx"), df.to_excel("seznam_najemov_new.xlsx"), doc1.save(prevzemni_name + ".docx"), convert("C:\\Users\\ikogovsek\\Desktop\\IGOR\\Python\\pythonProject\\" + pogodba_name + ".docx", "C:\\Users\\ikogovsek\\Desktop\\IGOR\\Python\\pythonProject\\PDF"), convert("C:\\Users\\ikogovsek\\Desktop\\IGOR\\Python\\pythonProject\\" + prevzemni_name + ".docx", "C:\\Users\\ikogovsek\\Desktop\\IGOR\\Python\\pythonProject\\PDF"), master3.destroy(), master.destroy()], width=10)
        yes_btn.pack(padx=10, pady=10, side=LEFT)
        no_btn = Button(frame1, text="Ne", bg="light blue", fg="red", command=master3.destroy, width=10)
        no_btn.pack(padx=10, pady=10, side=LEFT)

    else:
        master2 = tk.Tk()
        master2.title("SELMARapp - ERROR - manjkajo podatki")
        canvas = Canvas(master2, width=450, height=10)
        label = Label(master2, text="Opala!\n Niso izpolnjeni vsi podatki za kreiranje pogodbe in zapisnika.\n Prosim za izpolnitev vseh podatkov!", font=('Arial', 15))
        button = Button(master2, text="Ok", command=master2.destroy)
        button.grid(row=3, column=0)
        label.grid(row=1, column=0)
        canvas.grid()

        master2.mainloop()


master = tk.Tk()
master.title("SELMARapp - Najemna pogodba za pravno osebo")
tk.Label(master, text="Ime podjetja: ").grid(row=0)
tk.Label(master, text="Ulica: ").grid(row=1)
tk.Label(master, text="Hišna številka: ").grid(row=2)
tk.Label(master, text="Poštna številka: ").grid(row=3)
tk.Label(master, text="Mesto: ").grid(row=4)
tk.Label(master, text="Država: ").grid(row=5)
tk.Label(master, text="Matična številka: ").grid(row=6)
tk.Label(master, text="Davčna številka: ").grid(row=7)
tk.Label(master, text="Zastopnik: ").grid(row=8)
tk.Label(master, text="Ime in priimek zastopnika: ").grid(row=9)
tk.Label(master, text="Znamka: ").grid(row=10)
tk.Label(master, text="Model: ").grid(row=11)
tk.Label(master, text="VIN: ").grid(row=12)
tk.Label(master, text="Registrska oznaka (npr. CE AA-123): ").grid(row=13)
tk.Label(master, text="Registriran do: ").grid(row=14)
tk.Label(master, text="Najem od vključno: ").grid(row=15)
tk.Label(master, text="Najem do vključno: ").grid(row=16)
tk.Label(master, text="Obrok brez DDV: ").grid(row=17)
tk.Label(master, text="Cena na dan ali mesec: ").grid(row=18)
tk.Label(master, text="Kraj izdaje: ").grid(row=19)
tk.Label(master, text="Skrbnik: ").grid(row=20)


LOCATIONS = [
    "PE Celje",
    "PE Maribor",
    "PE Slovenska Bistrica"
]

PERSONS = [
    "Andrejčič Kevin",
    "Avberšek Sašo",
    "Bakračevič Aleksander",
    "Borko Borut",
    "Borko Robert",
    "Colja Danilo",
    "Golob Tadej",
    "Gostonj Srečko",
    "Kogovšek Igor",
    "Krajnc Klemen",
    "Kralj Boštjan",
    "Krumpak Tjaša",
    "Mašat Sandi",
    "Peklar Maja",
    "Sekulič Peter",
    "Seles Marcel"
]

FACTURING = [
    "mesec",
    "dan"
]

ZASTOPNIKI = [
    "direktor",
    "predsednik uprave",
    "prokurist",
    "ustanovitelj"

]

zastopnik = StringVar(master)
zastopnik.set("Izberi")

fakturiranje = StringVar(master)
fakturiranje.set("Izberi")

skrbnik = StringVar(master)
skrbnik.set("Izberi")

lokacija = StringVar(master)
lokacija.set("Izberi")

e1 = tk.Entry(master)
e2 = tk.Entry(master)
e3 = tk.Entry(master)
e4 = tk.Entry(master)
e5 = tk.Entry(master)
e6 = tk.Entry(master)
e7 = tk.Entry(master)
e8 = tk.Entry(master)
e9 = OptionMenu(master, zastopnik, *ZASTOPNIKI)
e10 = tk.Entry(master)
e11 = tk.Entry(master)
e12 = tk.Entry(master)
e13 = tk.Entry(master)
e14 = tk.Entry(master)
e15 = tk.Entry(master)
e16 = tk.Entry(master)
e17 = tk.Entry(master)
e18 = tk.Entry(master)
e19 = OptionMenu(master, fakturiranje, *FACTURING)
e20 = OptionMenu(master, lokacija, *LOCATIONS)
e21 = OptionMenu(master, skrbnik, *PERSONS)

e1.grid(row=0, column=1)
e2.grid(row=1, column=1)
e3.grid(row=2, column=1)
e4.grid(row=3, column=1)
e5.grid(row=4, column=1)
e6.grid(row=5, column=1)
e7.grid(row=6, column=1)
e8.grid(row=7, column=1)
e9.grid(row=8, column=1)
e10.grid(row=9, column=1)
e11.grid(row=10, column=1)
e12.grid(row=11, column=1)
e13.grid(row=12, column=1)
e14.grid(row=13, column=1)
e15.grid(row=14, column=1)
e16.grid(row=15, column=1)
e17.grid(row=16, column=1)
e18.grid(row=17, column=1)
e19.grid(row=18, column=1)
e20.grid(row=19, column=1)
e21.grid(row=20, column=1)

button1 = tk.Button(master,
          text='Nazaj',
          command=master.destroy).grid(row=21,
                                    column=0,
                                    sticky=tk.W,
                                    pady=4)
button2 = tk.Button(master,
          text='Kreiraj pogodbo in prevzemni zapisnik', command=kreiraj_pogodbo_in_zapisnik_PO).grid(row=21,
                                                                column=1,
                                                                sticky=tk.W,
                                                                pady=4)



master.mainloop()